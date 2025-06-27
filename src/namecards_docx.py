from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from .helpers import file_version_name

A4_h = Cm(29.7)
A4_w = Cm(21.0)

def create_namecard_docx(project):
    output_dir = Path(project.settings['output_folder'])

    docname = file_version_name(output_dir, 'namecards.docx')

    # Create a new document
    def all_png_files():
        return [Path(fn.name) for fn in output_dir.iterdir() 
                    if fn.suffix == '.png']
    doc = Document()

    # set margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.left_margin   = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.right_margin  = Cm(1.5)
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_height = A4_w
        sec.page_width = A4_h

    #place images in a table for formating 4 on each page
    tbl = doc.add_table(1, 2)
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(4.0)
    cells = tbl.rows[0].cells
    for i, png in enumerate(sorted(all_png_files(), key=lambda f: int(f.stem))):
        # format the cell and its paragraph
        cell = cells[i % 2]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.height = Cm(4.0)
        cell.width = Cm(4.0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format = p.paragraph_format
        format.space_before = Cm(0.0)
        format.space_after = Cm(0.3)
        # add new image
        pic = p.add_run().add_picture(str(output_dir / png), width=Pt(300))
        if (i+1) % 2 == 0: # add new row each 2nd cell
            cells = tbl.add_row().cells
        
    # Save the document
    save_path = output_dir / docname
    doc.save(save_path)

    return save_path
