from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from .tables import AllTables
from .persons import AllPersons

A4_h = Cm(29.7)
A4_w = Cm(21.0)

def create_special_foods_report(project):
    output_dir = Path(project.settings['output_folder'])

    doc = Document()

    # set margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.left_margin   = Cm(2.5)
        sec.bottom_margin = Cm(1.5)
        sec.right_margin  = Cm(2.5)
        sec.page_height = A4_h
        sec.page_width = A4_w

    doc.add_heading('Specialkost', 0)
    doc_tbl = doc.add_table(0, 3)

    persons = [p for p in sorted(project.persons.persons) 
               if p.special_foods]

    for i, p in enumerate(persons):
        row = doc_tbl.add_row()
        row.cells[0].width = Cm(1)
        row.cells[0].paragraphs[0].add_run(f'{i+1}')
        row.cells[1].paragraphs[0].add_run(f'{p.fname} {p.lname}').bold = True
        row.cells[2].paragraphs[0].add_run(f'{p.special_foods}').bold = True

    doc.add_page_break()

    doc.add_heading('Specialkost vid resp. bord')
    for tbl in project.tables.tables:
        if 0 == any(p.special_foods != '' for p in tbl.persons):
            continue
        heading = f'{tbl.id} \t{", ".join(tbl.departments())}'
        pr = doc.add_heading(heading, level=1)
        pr.add_run().bold = True
        doc_tbl = doc.add_table(0, 3)

        for i, p in enumerate(tbl.persons):
            if not p.special_foods:
                continue
            row = doc_tbl.add_row()
            row.cells[0].width = Cm(1)
            row.cells[0].paragraphs[0].add_run(f'{i+1}')
            row.cells[1].paragraphs[0].add_run(f'{p.fname} {p.lname}  ').bold = True
            row.cells[2].paragraphs[0].add_run(f'{p.special_foods}').italic = True

        
        doc.add_paragraph()

    doc.save(output_dir / 'special_foods.docx')