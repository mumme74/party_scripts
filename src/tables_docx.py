from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from .tables import AllTables

A4_h = Cm(29.7)
A4_w = Cm(21.0)

prj_dir = Path(__file__).parent.parent

def create_table_report(project):
    output_dir = Path(project.settings['output_folder'])
    template_path = Path(project.settings['table_sign']['file'])

     # Create a new document from template
    doc = Document(template_path)

    # set margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.left_margin   = Cm(2.5)
        sec.bottom_margin = Cm(1.5)
        sec.right_margin  = Cm(2.5)
        sec.page_height = A4_h
        sec.page_width = A4_w

    # one page for each table
    for i, tbl in enumerate(project.tables.tables):
        if i > 0:
            doc.add_page_break()
        doc.add_heading('Bordsplacering', 0)
        doc.add_paragraph()
        heading = f'{tbl.id} \t{", ".join(tbl.departments())}'
        pr = doc.add_heading(heading, level=1)
        pr.add_run().bold = True

        doc_tbl = doc.add_table(0, 3)

        for i, p in enumerate(tbl.persons):
            row = doc_tbl.add_row()
            row.cells[0].width = Cm(1)
            row.cells[0].paragraphs[0].add_run(f'{i+1}')
            row.cells[1].paragraphs[0].add_run(f'{p.fname} {p.lname}  ').bold = True
            row.cells[2].paragraphs[0].add_run(f'{p.dept.name}').italic = True

        doc.add_paragraph()
        doc.add_paragraph(f'Lediga platser: {tbl.free_seats()} av {tbl.num_seats}')

    doc.save(output_dir / 'table_placements.docx')